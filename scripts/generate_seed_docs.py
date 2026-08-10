"""Generates the synthetic seed documents committed under data/seed/.

These are the two documents the golden-dataset RAG evaluation harness
(tests/rag_eval/) ingests and asks questions against — see
tests/rag_eval/golden_dataset.py for the question set they back.

Not run automatically; the generated .docx files are committed to the
repo so the harness doesn't need python-docx generation as part of its
own test run. Re-run this script only if the seed content itself needs
to change:

    python scripts/generate_seed_docs.py

Section headings are written in ALL CAPS (not "1. Overview"-style
numbering) because app/ingestion/chunking.py's numbered-heading regex
requires a space directly after the number (e.g. "1 Overview"), not a
period ("1. Overview") — see PROJECT_PROGRESS.md's Phase 3 decision log.
ALL CAPS headings avoid that trap entirely and read naturally as SOP
section headers.

Body paragraphs are deliberately long (~1700-2000 characters each, most
of DEFAULT_TARGET_TOKENS=500's ~2000-character budget) so each section
lands in its own chunk rather than several short sections merging into
one — the golden dataset's expected_section_contains checks depend on
that one-section-per-chunk alignment holding for most sections.
"""

from __future__ import annotations

from pathlib import Path

import docx

SEED_DIR = Path(__file__).resolve().parent.parent / "data" / "seed"

PUMP_DOC = {
    "filename": "centrifugal_pump_p101_sop.docx",
    "title": "Centrifugal Pump P-101 Maintenance SOP",
    "sections": [
        (
            None,
            "Centrifugal Pump P-101 Maintenance SOP",
        ),
        (
            "OVERVIEW",
            "This standard operating procedure defines the maintenance requirements for "
            "centrifugal pump P-101, located in the process area of Plant 4. P-101 moves "
            "process fluid from the primary feed tank to the reactor inlet manifold at a "
            "nominal flow rate of 450 liters per minute. This document covers the four "
            "maintenance activities that keep P-101 in reliable service: bearing housing "
            "lubrication, vibration monitoring, mechanical seal replacement, and the safety "
            "precautions required before any technician performs work on the unit. Every "
            "activity described below has a defined interval or trigger condition, and every "
            "activity requires the pump to be properly isolated per the safety section before "
            "work begins. Deviating from these intervals without engineering sign-off is not "
            "permitted, since P-101 feeds a continuous reactor process and an unplanned "
            "failure here stops downstream production entirely.",
        ),
        (
            "LUBRICATION SCHEDULE",
            "The bearing housing on pump P-101 must be lubricated with ISO VG 32 turbine oil "
            "every 2,000 operating hours, tracked via the CMMS runtime counter rather than "
            "calendar time, since P-101 does not run continuously. Grease points at the "
            "outboard and inboard bearings should be serviced separately, every 500 operating "
            "hours, using a lithium complex grease rated for high-temperature service up to "
            "150 degrees Celsius. Apply grease slowly with a low-pressure grease gun and stop "
            "as soon as fresh grease is visible purging from the relief fitting; do not "
            "continue pumping grease past that point. Over-lubrication is one of the most "
            "common causes of premature bearing failure on pumps of this class, because excess "
            "grease raises internal case pressure and churns, generating heat that breaks down "
            "the lubricant film faster than normal operation would. Under-lubrication is "
            "equally damaging, so if a scheduled interval is missed, do not simply skip it; log "
            "the deviation and service the bearing at the next available opportunity rather "
            "than waiting for the following scheduled interval.",
        ),
        (
            "VIBRATION MONITORING",
            "Vibration readings must be collected monthly at the drive-end and non-drive-end "
            "bearing housings using a calibrated handheld vibration analyzer, with data logged "
            "against the equipment ID in the CMMS so trends are visible over time rather than "
            "just the latest reading. The alarm threshold for P-101 is 7.1 millimeters per "
            "second RMS overall velocity, measured per ISO 10816 guidance for this pump class "
            "and mounting configuration. Any single reading exceeding this threshold requires "
            "an immediate work order and a root-cause investigation before the pump is returned "
            "to continuous service; do not simply re-run the reading and accept a lower repeat "
            "value as clearing the alarm. A reading that is trending upward across three or "
            "more consecutive monthly checks, even if still under the 7.1 mm/s threshold, "
            "should also be flagged for review, since a rising trend is often the earliest "
            "available warning of bearing wear or developing misalignment, well before an "
            "absolute threshold would be crossed.",
        ),
        (
            "SEAL REPLACEMENT",
            "The mechanical seal on P-101 should be replaced on a preventive basis every 18 "
            "months, or immediately upon any visible leakage at the seal faces, whichever comes "
            "first. Preventive replacement on schedule is strongly preferred over run-to-failure "
            "for this seal, because a failed mechanical seal on P-101 releases process fluid "
            "directly into the pump pedestal area rather than into a contained drain, which "
            "turns what would otherwise be a routine planned replacement into an unplanned "
            "cleanup and containment event. Replacement requires isolating and fully draining "
            "the pump casing before the seal cartridge is removed; refer to the pump OEM manual "
            "for the exact seal cartridge torque specification, since over-torquing a new seal "
            "during installation is a common cause of early failure of the replacement seal "
            "itself, sometimes within days of being installed.",
        ),
        (
            "SAFETY PRECAUTIONS",
            "Lockout/tagout procedures must be applied and independently verified before any "
            "maintenance task on P-101, including tasks that might seem minor, such as routine "
            "lubrication or taking a vibration reading with the pump stopped. Verification means "
            "a second qualified person confirms zero energy state at the pump before work "
            "begins, not just that the lockout device has been physically applied. Do not remove "
            "lockout devices until every technician who applied a personal lock has confirmed "
            "their portion of the work is complete and the immediate work area has been visually "
            "cleared of tools and loose material. P-101 handles process fluid that is mildly "
            "corrosive; technicians performing seal replacement or any task that could expose "
            "them to residual fluid in the casing must wear the chemical-resistant gloves and "
            "face shield specified in the site's chemical handling procedure for this fluid "
            "class, in addition to standard PPE.",
        ),
    ],
}

CONVEYOR_DOC = {
    "filename": "conveyor_cb200_manual.docx",
    "title": "Conveyor Belt CB-200 Operations and Maintenance Manual",
    "sections": [
        (
            None,
            "Conveyor Belt CB-200 Operations and Maintenance Manual",
        ),
        (
            "OVERVIEW",
            "This manual covers routine operation and maintenance for conveyor CB-200, which "
            "transports finished goods from the end of the packaging line to the warehouse "
            "staging area over a span of approximately 60 meters. CB-200 runs on a fixed "
            "schedule during production shifts and is normally stopped, but not de-energized, "
            "between shifts. This document covers the four maintenance activities that keep "
            "CB-200 running reliably: belt tensioning, drive motor inspection, emergency stop "
            "testing, and troubleshooting steps for the most common fault this conveyor "
            "experiences, belt tracking drift. Anyone performing maintenance on CB-200 should "
            "read the relevant section fully before starting work, since several of the checks "
            "described below (particularly belt tensioning) are only meaningful if performed in "
            "the specific sequence described.",
        ),
        (
            "BELT TENSIONING",
            "Belt tension on CB-200 must be checked weekly, with the conveyor stopped and "
            "unloaded, using the tension gauge mounted at the take-up frame near the tail "
            "pulley. The target sag at mid-span is 1 to 2 percent of the total span length, "
            "measured as the vertical deflection at the center point when a fixed test weight "
            "is applied. Excessive sag beyond this range causes the belt to slip against the "
            "drive pulley under load, which both wastes energy and accelerates wear on the "
            "pulley lagging. Insufficient sag, meaning the belt is over-tensioned, accelerates "
            "bearing wear on the take-up pulley and increases motor load even when the conveyor "
            "is running empty. Adjust tension in small increments at the take-up screws, no more "
            "than a quarter turn at a time on each side, and re-measure sag after every "
            "adjustment rather than making several turns and checking once at the end, since "
            "over-adjustment is easy to introduce and hard to diagnose after the fact. Record "
            "the measured sag value and the date of each weekly check in the CMMS, since a "
            "sudden change between two consecutive weekly checks is often a more useful early "
            "warning sign than the absolute sag value alone, particularly for catching a "
            "developing idler or pulley problem before it becomes visible tracking drift.",
        ),
        (
            "MOTOR INSPECTION",
            "The drive motor bearings on CB-200 should be inspected quarterly for unusual "
            "noise, vibration, or excessive heat at the bearing housings, using a handheld "
            "temperature probe and a simple listen-and-feel check with the conveyor running "
            "under normal load. Motor bearings on CB-200 are scheduled for preventive "
            "replacement at 20,000 operating hours regardless of their condition at that point, "
            "as a deliberate preventive measure against unplanned downtime on a conveyor that "
            "sits directly downstream of the packaging line with no buffer capacity. A bearing "
            "that fails without warning here does not just stop CB-200, it backs up the entire "
            "packaging line within minutes. If a quarterly inspection finds any bearing running "
            "noticeably hotter than the others, or producing an audible grinding or knocking "
            "sound rather than a smooth hum, schedule replacement immediately rather than "
            "waiting for the next quarterly check, even if the 20,000 hour interval has not yet "
            "been reached.",
        ),
        (
            "EMERGENCY STOP TESTING",
            "Emergency stop pull cords running along the full length of CB-200 must be tested "
            "monthly with the conveyor running under normal, unloaded conditions. Pull the cord "
            "firmly at three points along the conveyor for each test: the head end near the "
            "discharge, mid-span, and the tail end near the infeed, and confirm at each point "
            "that the conveyor comes to a complete stop within the rated stopping distance "
            "specified on the equipment nameplate. After each test, the pull cord switch must "
            "be manually reset at its physical location before the conveyor can be restarted; "
            "if a switch does not reset cleanly or requires unusual force, treat this as a "
            "failed test and raise a work order rather than forcing the reset and returning the "
            "conveyor to service. A documented log of monthly e-stop tests, including the date, "
            "tester, and stopping distance observed at each of the three points, must be kept "
            "for this conveyor.",
        ),
        (
            "TROUBLESHOOTING",
            "Belt tracking issues, meaning the belt gradually drifting toward one side of the "
            "conveyor frame during operation, are the most common fault reported on CB-200. The "
            "most frequent root cause is uneven belt tension side to side, so the first "
            "troubleshooting step is always to re-check tension per the tensioning section "
            "above before adjusting anything else. If tension is confirmed even on both sides "
            "and drift continues, inspect the idler rollers along the drift side for "
            "misalignment or a seized bearing that is causing uneven rolling resistance. "
            "Material buildup on the drive or tail pulley face is a third, less common cause; "
            "buildup changes the effective pulley diameter unevenly across its width and can "
            "steer the belt just as effectively as a physical misalignment. Only after tension, "
            "idler alignment, and pulley cleanliness have all been checked and ruled out should "
            "the tracking guide adjusters themselves be used to correct drift, since adjusting "
            "guides to mask an underlying tension or alignment problem tends to make the "
            "underlying problem worse over time.",
        ),
    ],
}


def _write_doc(spec: dict) -> Path:
    document = docx.Document()
    for heading, body in spec["sections"]:
        if heading is not None:
            document.add_paragraph(heading)
        document.add_paragraph(body)

    SEED_DIR.mkdir(parents=True, exist_ok=True)
    path = SEED_DIR / spec["filename"]
    document.save(str(path))
    return path


def main() -> None:
    for spec in (PUMP_DOC, CONVEYOR_DOC):
        path = _write_doc(spec)
        print(f"wrote {path}")


if __name__ == "__main__":
    main()
